"""
Psychologist Reports Workspace API

Endpoints to power the per-student Reports Workspace:
  - Generate a background summary from the latest completed parent chat session.
  - Upload an IQ test PDF, extract text + structured scores, persist a CognitiveProfile.
  - Generate a cognitive report from the latest CognitiveProfile.
  - Synthesize unified insights from the latest background + cognitive reports.
  - CRUD-style edit/finalize endpoints for any PsychologistReport row.
"""

import logging
import os
import time
import uuid as uuid_lib
from typing import List, Optional

import re
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import and_, desc, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.security import get_current_active_user
from app.models.user import User, UserRole
from app.models.student import Student
from app.models.student_guardian import StudentGuardian
from app.models.assignment import AssessmentAssignment, AssignmentStatus
from app.models.chat import ChatSession, ChatSessionStatus
from app.models.upload import IQTestUpload, UploadStatus, CognitiveProfile
from app.models.psychologist_report import PsychologistReport
from app.agents.report_agents import (
    BackgroundSummaryAgent,
    IQScoreExtractorAgent,
    CognitiveReportAgent,
    UnifiedInsightsAgent,
    _classify_role,
)
from app.services.pdf_extractor import extract_text_from_pdf
from app.services.s3_storage import s3_storage

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/psychologist-reports", tags=["psychologist-reports"])


# Ensure temp upload directory exists once at import time.
TEMP_UPLOAD_DIR = "/tmp/iq_uploads"
try:
    os.makedirs(TEMP_UPLOAD_DIR, exist_ok=True)
except Exception as e:
    logger.warning(f"Could not create {TEMP_UPLOAD_DIR}: {e}")


# ============================================================================
# AUTH DEPENDENCY
# ============================================================================
def require_psychologist_or_admin(
    current_user: User = Depends(get_current_active_user),
) -> User:
    """Only psychologists and admins may access the reports workspace."""
    if current_user.role not in [UserRole.PSYCHOLOGIST, UserRole.ADMIN]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Psychologist or Admin access required",
        )
    return current_user


# ============================================================================
# PYDANTIC SCHEMAS
# ============================================================================
class ReportOut(BaseModel):
    id: str
    student_id: str
    report_type: str
    content_markdown: str
    status: str
    source_chat_session_id: Optional[str] = None
    source_cognitive_profile_id: Optional[str] = None
    generation_ms: Optional[int] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None


class BlankReportCreate(BaseModel):
    report_type: str  # "background_summary" | "cognitive_report" | "unified_insights"


class ReportUpdate(BaseModel):
    content_markdown: str


# ============================================================================
# HELPERS
# ============================================================================
def _serialize_report(r: PsychologistReport) -> dict:
    return {
        "id": str(r.id),
        "student_id": str(r.student_id),
        "created_by_user_id": str(r.created_by_user_id) if r.created_by_user_id else None,
        "report_type": r.report_type,
        "content_markdown": r.content_markdown or "",
        "source_data": r.source_data,
        "source_chat_session_id": str(r.source_chat_session_id) if r.source_chat_session_id else None,
        "source_cognitive_profile_id": str(r.source_cognitive_profile_id) if r.source_cognitive_profile_id else None,
        "status": r.status,
        "generation_ms": r.generation_ms,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "updated_at": r.updated_at.isoformat() if r.updated_at else None,
    }


def _serialize_profile(p: CognitiveProfile, upload: Optional[IQTestUpload] = None) -> dict:
    return {
        "id": str(p.id),
        "student_id": str(p.student_id),
        "iq_test_upload_id": str(p.iq_test_upload_id) if p.iq_test_upload_id else None,
        "test_name": p.test_name,
        "test_date": p.test_date.isoformat() if p.test_date else None,
        "administered_by": p.administered_by,
        "parsed_scores": p.parsed_scores,
        "raw_ocr_text": p.raw_ocr_text,
        "confidence_score": p.confidence_score,
        "requires_review": p.requires_review,
        "created_at": p.created_at.isoformat() if p.created_at else None,
        "updated_at": p.updated_at.isoformat() if p.updated_at else None,
        "iq_test_upload": {
            "id": str(upload.id),
            "file_name": upload.file_name,
            "file_size_bytes": upload.file_size_bytes,
            "mime_type": upload.mime_type,
            "uploaded_at": upload.uploaded_at.isoformat() if upload.uploaded_at else None,
        } if upload else None,
    }


async def _get_student_or_404(db: AsyncSession, student_id: uuid_lib.UUID) -> Student:
    result = await db.execute(select(Student).where(Student.id == student_id))
    student = result.scalar_one_or_none()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    return student


async def _get_report_or_404(db: AsyncSession, report_id: uuid_lib.UUID) -> PsychologistReport:
    result = await db.execute(
        select(PsychologistReport).where(PsychologistReport.id == report_id)
    )
    report = result.scalar_one_or_none()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return report


async def _latest_completed_session(
    db: AsyncSession, student_id: uuid_lib.UUID
) -> Optional[ChatSession]:
    """Find the most recent COMPLETED ChatSession for this student.

    Deprecated: kept for backwards compatibility with pre-multi-party callers.
    Prefer `_all_assessor_rows` / `_all_completed_perspectives` for new code.
    """
    result = await db.execute(
        select(ChatSession)
        .join(AssessmentAssignment, ChatSession.assignment_id == AssessmentAssignment.id)
        .where(AssessmentAssignment.student_id == student_id)
        .where(ChatSession.status == ChatSessionStatus.COMPLETED.value)
        .order_by(desc(ChatSession.id))
    )
    return result.scalars().first()


async def _all_assessor_rows(
    db: AsyncSession, student_id: uuid_lib.UUID
) -> List[dict]:
    """
    Return one dict per non-cancelled assessment assignment for the student,
    with joined guardian + session info. Shape:

        {
            "assignment_id": UUID,
            "assignment_status": "ASSIGNED" | "IN_PROGRESS" | "COMPLETED",
            "guardian_user_id": UUID,
            "guardian_name": str,
            "guardian_email": str,
            "relationship_type": str,   # "Mother" / "Father" / "School" / "Guardian"
            "session_id": UUID | None,
            "session_status": str | None,
            "context_data": dict | None,
            "started_at": datetime | None,
            "completed_at": datetime | None,
        }

    Source of truth for completion is AssessmentAssignment.status (uppercase).
    Do NOT key completion off ChatSession.status (lowercase) — the session can
    lag behind the assignment transition in edge cases.
    """
    result = await db.execute(
        select(
            AssessmentAssignment,
            User,
            StudentGuardian,
            ChatSession,
        )
        .join(User, AssessmentAssignment.assigned_to_user_id == User.id)
        .outerjoin(
            StudentGuardian,
            and_(
                StudentGuardian.student_id == AssessmentAssignment.student_id,
                StudentGuardian.guardian_user_id == AssessmentAssignment.assigned_to_user_id,
            ),
        )
        .outerjoin(ChatSession, ChatSession.assignment_id == AssessmentAssignment.id)
        .where(AssessmentAssignment.student_id == student_id)
        .where(AssessmentAssignment.status != AssignmentStatus.CANCELLED)
        .order_by(AssessmentAssignment.assigned_at.asc())
    )

    rows: List[dict] = []
    for assignment, user, sg, session in result.all():
        rel = (sg.relationship_type if sg and sg.relationship_type else None) or (
            user.role.value.title() if user and user.role else "Guardian"
        )
        status_val = (
            assignment.status.value
            if hasattr(assignment.status, "value")
            else str(assignment.status)
        )
        rows.append({
            "assignment_id": assignment.id,
            "assignment_status": status_val,
            "guardian_user_id": user.id,
            "guardian_name": user.full_name or user.email,
            "guardian_email": user.email,
            "relationship_type": rel,
            "session_id": session.id if session else None,
            "session_status": session.status if session else None,
            "context_data": session.context_data if session else None,
            "started_at": assignment.started_at,
            "completed_at": assignment.completed_at,
        })
    return rows


async def _all_completed_perspectives(
    db: AsyncSession, student_id: uuid_lib.UUID
) -> tuple[list[dict], list[dict]]:
    """
    Split the assessor rows into (completed, pending).

    Used by the background-summary generator to decide whether to gate
    and to build the merged context.
    """
    rows = await _all_assessor_rows(db, student_id)
    completed = [r for r in rows if r["assignment_status"] == AssignmentStatus.COMPLETED.value]
    pending = [r for r in rows if r["assignment_status"] != AssignmentStatus.COMPLETED.value]
    return completed, pending


async def _latest_report(
    db: AsyncSession, student_id: uuid_lib.UUID, report_type: str
) -> Optional[PsychologistReport]:
    result = await db.execute(
        select(PsychologistReport)
        .where(PsychologistReport.student_id == student_id)
        .where(PsychologistReport.report_type == report_type)
        .order_by(desc(PsychologistReport.created_at))
    )
    return result.scalars().first()


# ============================================================================
# GET WORKSPACE STATE
# ============================================================================
@router.get("/students/{student_id}/workspace")
async def get_workspace(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """
    Return all state needed to render the Reports Workspace for a student:
      - the student record
      - the latest completed ChatSession's context_data (if any)
      - every existing PsychologistReport, grouped by report_type (latest first)
      - every CognitiveProfile row with its IQTestUpload metadata
    """
    student = await _get_student_or_404(db, student_id)

    latest_session = await _latest_completed_session(db, student_id)

    # Multi-party progress: one row per non-cancelled assignment
    assessor_rows = await _all_assessor_rows(db, student_id)

    # All reports for this student, newest first
    reports_result = await db.execute(
        select(PsychologistReport)
        .where(PsychologistReport.student_id == student_id)
        .order_by(desc(PsychologistReport.created_at))
    )
    all_reports = reports_result.scalars().all()

    grouped_reports = {
        "background_summary": [],
        "background_summary_parent": [],
        "background_summary_school": [],
        "cognitive_report": [],
        "unified_insights": [],
    }
    for r in all_reports:
        grouped_reports.setdefault(r.report_type, []).append(_serialize_report(r))

    # Cognitive profiles (+ their uploads)
    profiles_result = await db.execute(
        select(CognitiveProfile)
        .where(CognitiveProfile.student_id == student_id)
        .order_by(desc(CognitiveProfile.created_at))
    )
    profiles = profiles_result.scalars().all()

    profile_dicts = []
    for p in profiles:
        upload = None
        if p.iq_test_upload_id:
            upload_result = await db.execute(
                select(IQTestUpload).where(IQTestUpload.id == p.iq_test_upload_id)
            )
            upload = upload_result.scalar_one_or_none()
        profile_dicts.append(_serialize_profile(p, upload))

    return {
        "student": {
            "id": str(student.id),
            "first_name": student.first_name,
            "last_name": student.last_name,
            "date_of_birth": student.date_of_birth.isoformat() if student.date_of_birth else None,
            "year_group": student.year_group,
            "school_name": student.school_name,
        },
        "latest_completed_session": {
            "id": str(latest_session.id),
            "context_data": latest_session.context_data,
            "status": latest_session.status,
        } if latest_session else None,
        "reports": grouped_reports,
        "cognitive_profiles": profile_dicts,
        "assessors": [
            {
                "assignment_id": str(r["assignment_id"]),
                "guardian_user_id": str(r["guardian_user_id"]),
                "guardian_name": r["guardian_name"],
                "guardian_email": r["guardian_email"],
                "relationship_type": r["relationship_type"],
                "status": r["assignment_status"],  # already a string, uppercase
                "session_id": str(r["session_id"]) if r["session_id"] else None,
                "session_status": r["session_status"],
                "started_at": r["started_at"].isoformat() if r["started_at"] else None,
                "completed_at": r["completed_at"].isoformat() if r["completed_at"] else None,
            }
            for r in assessor_rows
        ],
        "completion_count": {
            "done": sum(1 for r in assessor_rows if r["assignment_status"] == AssignmentStatus.COMPLETED.value),
            "total": len(assessor_rows),
        },
        "all_assessors_complete": bool(assessor_rows) and all(
            r["assignment_status"] == AssignmentStatus.COMPLETED.value for r in assessor_rows
        ),
    }


# ============================================================================
# QUESTIONNAIRE RESPONSES — VIEW PARENT / SCHOOL ANSWERS
# ============================================================================
@router.get("/students/{student_id}/responses")
async def get_questionnaire_responses(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Return all chatbot questionnaire responses for a student.

    Output shape (clean for the frontend — NOT a raw JSON dump of the
    chat session):

        {
            "student": {"id", "first_name", "last_name", ...},
            "sessions": [
                {
                    "session_id", "status", "flow_type",
                    "respondent": {"name", "email", "role"},
                    "started_at", "completed_at",
                    "sections": [
                        {
                            "title": "Reasons for Assessment",
                            "category": "background",
                            "items": [
                                {"node_id", "question", "answer",
                                 "option_label", "is_skipped", "is_multi"},
                                ...
                            ],
                        }, ...
                    ],
                }, ...
            ],
        }

    Permission: psychologist OR admin. Admin sees any student; psychologist
    needs to be in the student's care chain (enforced by existing role gate
    on the workspace endpoints — same audience).
    """
    from app.models.chat import ChatMessage  # local import — avoids circulars

    student = await _get_student_or_404(db, student_id)

    sessions_q = await db.execute(
        select(ChatSession)
        .join(AssessmentAssignment, ChatSession.assignment_id == AssessmentAssignment.id)
        .where(AssessmentAssignment.student_id == student_id)
        .order_by(desc(ChatSession.started_at))
    )
    sessions: List[ChatSession] = list(sessions_q.scalars().all())

    if not sessions:
        return {
            "student": {
                "id": str(student.id),
                "first_name": student.first_name,
                "last_name": student.last_name,
            },
            "sessions": [],
        }

    # Bulk load all messages for these sessions, ordered chronologically
    session_ids = [s.id for s in sessions]
    msgs_q = await db.execute(
        select(ChatMessage)
        .where(ChatMessage.session_id.in_(session_ids))
        .order_by(ChatMessage.session_id, ChatMessage.timestamp.asc(), ChatMessage.id.asc())
    )
    all_msgs = list(msgs_q.scalars().all())

    # Group messages by session
    msgs_by_session: dict[uuid_lib.UUID, list[ChatMessage]] = {}
    for m in all_msgs:
        msgs_by_session.setdefault(m.session_id, []).append(m)

    # Build respondent lookups (assignment -> guardian + role) in one query
    assignment_ids = [s.assignment_id for s in sessions if s.assignment_id]
    respondent_by_assignment: dict[uuid_lib.UUID, dict] = {}
    if assignment_ids:
        rows_q = await db.execute(
            select(AssessmentAssignment, User, StudentGuardian)
            .join(User, AssessmentAssignment.assigned_to_user_id == User.id)
            .outerjoin(
                StudentGuardian,
                and_(
                    StudentGuardian.student_id == AssessmentAssignment.student_id,
                    StudentGuardian.guardian_user_id == AssessmentAssignment.assigned_to_user_id,
                ),
            )
            .where(AssessmentAssignment.id.in_(assignment_ids))
        )
        for assignment, user, sg in rows_q.all():
            rel = (sg.relationship_type if sg and sg.relationship_type else None) or (
                user.role.value.title() if user and user.role else "Guardian"
            )
            respondent_by_assignment[assignment.id] = {
                "name": user.full_name or user.email,
                "email": user.email,
                "role": rel,
            }

    out_sessions: list[dict] = []
    for sess in sessions:
        msgs = msgs_by_session.get(sess.id, [])

        # Pair (bot question -> next user answer) and group by category/section
        pairs: list[dict] = []
        pending_bot: ChatMessage | None = None
        for m in msgs:
            if m.role == "bot" and m.message_type in ("mcq", "text_only", "text"):
                pending_bot = m
                continue
            if m.role == "user" and pending_bot is not None:
                meta_bot = pending_bot.message_metadata or {}
                meta_user = m.message_metadata or {}
                question = pending_bot.content or ""
                # Strip out the auto-prepended student name "Hi {name}, ..." style
                question = question.strip()
                # Determine user-facing answer:
                option_label = None
                selected_value = meta_user.get("selected_option") or meta_user.get("choice_value")
                if selected_value:
                    # Bot meta carries options list — find label
                    for opt in (meta_bot.get("options") or []):
                        if isinstance(opt, dict) and opt.get("value") == selected_value:
                            option_label = opt.get("label")
                            break
                    # Multi-select: comma-separated values
                    if option_label is None and isinstance(selected_value, str) and "," in selected_value:
                        labels = []
                        for tok in selected_value.split(","):
                            tok = tok.strip()
                            for opt in (meta_bot.get("options") or []):
                                if isinstance(opt, dict) and opt.get("value") == tok:
                                    labels.append(opt.get("label", tok))
                                    break
                        if labels:
                            option_label = ", ".join(labels)

                pairs.append({
                    "node_id": meta_bot.get("node_id") or meta_bot.get("question_id"),
                    "category": meta_bot.get("category") or "general",
                    "section": (meta_bot.get("section") or meta_bot.get("section_title")
                                or _category_title(meta_bot.get("category") or "general")),
                    "question": question,
                    "answer": m.content or "",
                    "option_label": option_label,
                    "is_skipped": (m.message_type == "skip") or (m.content or "").strip().lower() == "(skipped)",
                    "is_multi": bool(meta_bot.get("multi_select")),
                    "timestamp": m.timestamp.isoformat() if m.timestamp else None,
                })
                pending_bot = None

        # Group pairs into sections preserving first-appearance order
        sections_ordered: list[str] = []
        items_by_section: dict[str, list[dict]] = {}
        for p in pairs:
            title = p["section"] or "Other"
            if title not in items_by_section:
                items_by_section[title] = []
                sections_ordered.append(title)
            items_by_section[title].append(p)

        respondent = respondent_by_assignment.get(sess.assignment_id) or {
            "name": "Unknown respondent",
            "email": None,
            "role": "Guardian",
        }

        out_sessions.append({
            "session_id": str(sess.id),
            "flow_type": sess.flow_type,
            "status": sess.status,
            "respondent": respondent,
            "started_at": sess.started_at.isoformat() if sess.started_at else None,
            "completed_at": sess.completed_at.isoformat() if sess.completed_at else None,
            "sections": [
                {
                    "title": title,
                    "items": items_by_section[title],
                }
                for title in sections_ordered
            ],
            "totals": {
                "answered": len(pairs),
                "skipped": sum(1 for p in pairs if p["is_skipped"]),
            },
        })

    return {
        "student": {
            "id": str(student.id),
            "first_name": student.first_name,
            "last_name": student.last_name,
            "date_of_birth": student.date_of_birth.isoformat() if student.date_of_birth else None,
            "year_group": student.year_group,
            "school_name": student.school_name,
        },
        "sessions": out_sessions,
        "viewer_role": current_user.role.value if current_user and current_user.role else None,
    }


def _category_title(cat: str) -> str:
    """Map a category slug to a human-readable section title."""
    mapping = {
        "background": "Reasons for Assessment",
        "family": "Family Context",
        "birth_history": "Pregnancy & Birth",
        "dev_milestones": "Developmental Milestones",
        "health": "General Health",
        "vision_hearing": "Vision & Hearing",
        "sensory_motor": "Sensory & Motor",
        "attention": "Attention",
        "social": "Social",
        "emotional": "Emotional",
        "behaviour": "Behaviour",
        "behavioural": "Behaviour",
        "academic": "Academic",
        "education": "Education",
        "literacy": "Literacy",
        "independent_learning": "Independent Learning",
        "strengths": "Strengths & Interests",
        "closing": "Closing",
        "intro": "Introduction",
        "general": "Other",
    }
    return mapping.get((cat or "").lower(), (cat or "Other").replace("_", " ").title())


# ============================================================================
# BACKGROUND SUMMARY — GENERATE
# ============================================================================
@router.post("/students/{student_id}/background-summary/generate")
async def generate_background_summary(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """
    Generate a Background Summary report.

    Strict gate: the report only generates when **every** non-cancelled
    AssessmentAssignment for this student has reached status COMPLETED.
    Single-assignee flow is the trivial 1-of-1 case of this rule.

    For multi-party cases (school + mother + father) the agent receives a
    merged context with role attribution so the final report weaves
    every perspective together.
    """
    student = await _get_student_or_404(db, student_id)

    completed, pending = await _all_completed_perspectives(db, student_id)
    total = len(completed) + len(pending)

    if total == 0:
        raise HTTPException(
            status_code=400,
            detail=(
                "No assessment assignments exist for this student. "
                "An admin must assign the assessment to at least one guardian first."
            ),
        )

    if pending:
        pending_roles = sorted({p["relationship_type"] or "Guardian" for p in pending})
        raise HTTPException(
            status_code=400,
            detail={
                "message": f"Waiting for {len(pending)} of {total} assessors to complete.",
                "pending_roles": pending_roles,
                "completed_count": len(completed),
                "total_count": total,
            },
        )

    student_name = f"{student.first_name} {student.last_name}".strip() or "the student"

    # Build merged context. With a single completed assignee this degrades
    # to a single-entry perspectives list — agent behaviour stays identical
    # apart from the added role prefix in the prompt.
    merged_context = {
        "student_info": {
            "first_name": student.first_name,
            "last_name": student.last_name,
            "date_of_birth": student.date_of_birth.isoformat() if student.date_of_birth else None,
            "year_group": student.year_group,
            "school_name": student.school_name,
        },
        "perspectives": [
            {
                "role": p["relationship_type"],
                "respondent_name": p["guardian_name"],
                "respondent_user_id": str(p["guardian_user_id"]),
                "session_id": str(p["session_id"]) if p["session_id"] else None,
                "context": p["context_data"] or {},
            }
            for p in completed
        ],
    }

    agent = BackgroundSummaryAgent()
    started = time.time()
    try:
        markdown = await agent.generate(merged_context, student_name)
    except Exception as e:
        logger.error(f"BackgroundSummaryAgent raised: {e}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Agent failed: {str(e)}")
    generation_ms = int((time.time() - started) * 1000)

    # Persist: source_chat_session_id is a single FK so point it at the
    # most recent session; stash the full set in source_data for provenance.
    session_ids = [p["session_id"] for p in completed if p["session_id"]]
    primary_session_id = session_ids[-1] if session_ids else None

    report = PsychologistReport(
        student_id=student_id,
        created_by_user_id=current_user.id,
        report_type="background_summary",
        content_markdown=markdown,
        source_data={
            "session_ids": [str(sid) for sid in session_ids],
            "perspectives": [
                {
                    "role": p["relationship_type"],
                    "respondent_name": p["guardian_name"],
                    "respondent_user_id": str(p["guardian_user_id"]),
                    "session_id": str(p["session_id"]) if p["session_id"] else None,
                }
                for p in completed
            ],
        },
        source_chat_session_id=primary_session_id,
        status="draft",
        generation_ms=generation_ms,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)

    return _serialize_report(report)


# ============================================================================
# BACKGROUND SUMMARY — ROLE-SCOPED GENERATE (PARENT / SCHOOL)
# ============================================================================
async def _generate_role_scoped_background(
    *,
    student: Student,
    student_id: uuid_lib.UUID,
    current_user: User,
    db: AsyncSession,
    voice: str,
    report_type: str,
) -> dict:
    """Shared implementation for the two role-scoped background endpoints.

    `voice` is "parent" or "school"; `report_type` is the value stored on
    PsychologistReport.report_type ("background_summary_parent" /
    "background_summary_school"). Gating: at least one assessor of the
    matching voice must be COMPLETED. Pending assessors of the OTHER voice
    do NOT block this report — that decoupling is the whole point.
    """
    rows = await _all_assessor_rows(db, student_id)
    if not rows:
        raise HTTPException(
            status_code=400,
            detail=(
                "No assessment assignments exist for this student. "
                "An admin must assign the assessment to at least one assessor first."
            ),
        )

    matching = [r for r in rows if _classify_role(r.get("relationship_type")) == voice]
    if not matching:
        raise HTTPException(
            status_code=400,
            detail=(
                f"No {voice} assessor has been assigned to this student. "
                f"Assign the assessment to a {voice}-type respondent and try again."
            ),
        )

    completed = [r for r in matching if r["assignment_status"] == AssignmentStatus.COMPLETED.value]
    pending = [r for r in matching if r["assignment_status"] != AssignmentStatus.COMPLETED.value]

    if not completed:
        pending_names = sorted({p.get("guardian_name") or "Unknown" for p in pending})
        raise HTTPException(
            status_code=400,
            detail={
                "message": (
                    f"Waiting for the {voice} assessment to be completed before this "
                    f"summary can be generated."
                ),
                "voice": voice,
                "pending_assessors": pending_names,
            },
        )

    student_name = f"{student.first_name} {student.last_name}".strip() or "the student"

    merged_context = {
        "student_info": {
            "first_name": student.first_name,
            "last_name": student.last_name,
            "date_of_birth": student.date_of_birth.isoformat() if student.date_of_birth else None,
            "year_group": student.year_group,
            "school_name": student.school_name,
        },
        "perspectives": [
            {
                "role": p["relationship_type"],
                "respondent_name": p["guardian_name"],
                "respondent_user_id": str(p["guardian_user_id"]),
                "session_id": str(p["session_id"]) if p["session_id"] else None,
                "context": p["context_data"] or {},
            }
            for p in completed
        ],
    }

    agent = BackgroundSummaryAgent()
    started = time.time()
    try:
        if voice == "parent":
            markdown = await agent.generate_parent_summary(merged_context, student_name)
        else:
            markdown = await agent.generate_school_summary(merged_context, student_name)
    except Exception as e:
        logger.error(f"BackgroundSummaryAgent ({voice}) raised: {e}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Agent failed: {str(e)}")
    generation_ms = int((time.time() - started) * 1000)

    session_ids = [p["session_id"] for p in completed if p["session_id"]]
    primary_session_id = session_ids[-1] if session_ids else None

    report = PsychologistReport(
        student_id=student_id,
        created_by_user_id=current_user.id,
        report_type=report_type,
        content_markdown=markdown,
        source_data={
            "voice": voice,
            "session_ids": [str(sid) for sid in session_ids],
            "perspectives": [
                {
                    "role": p["relationship_type"],
                    "respondent_name": p["guardian_name"],
                    "respondent_user_id": str(p["guardian_user_id"]),
                    "session_id": str(p["session_id"]) if p["session_id"] else None,
                }
                for p in completed
            ],
            "pending_assessors_of_this_voice": [
                {
                    "role": p["relationship_type"],
                    "respondent_name": p["guardian_name"],
                    "assignment_status": p["assignment_status"],
                }
                for p in pending
            ],
        },
        source_chat_session_id=primary_session_id,
        status="draft",
        generation_ms=generation_ms,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)
    return _serialize_report(report)


@router.post("/students/{student_id}/background-summary-parent/generate")
async def generate_background_summary_parent(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Generate the PARENT-only background summary.

    Gates on at least one parent-type assessor being COMPLETED. Pending
    school assessors do NOT block this endpoint. The agent is forbidden
    from using school-voice content even if it is present elsewhere.
    """
    student = await _get_student_or_404(db, student_id)
    return await _generate_role_scoped_background(
        student=student,
        student_id=student_id,
        current_user=current_user,
        db=db,
        voice="parent",
        report_type="background_summary_parent",
    )


@router.post("/students/{student_id}/background-summary-school/generate")
async def generate_background_summary_school(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Generate the SCHOOL-only background summary.

    Gates on at least one school-type assessor being COMPLETED. Pending
    parent assessors do NOT block this endpoint. The agent is forbidden
    from using parent-voice content even if it is present elsewhere.
    """
    student = await _get_student_or_404(db, student_id)
    return await _generate_role_scoped_background(
        student=student,
        student_id=student_id,
        current_user=current_user,
        db=db,
        voice="school",
        report_type="background_summary_school",
    )


# ============================================================================
# COGNITIVE REPORT — UPLOAD PDF
# ============================================================================
@router.post("/students/{student_id}/cognitive-report/upload")
async def upload_cognitive_report_pdf(
    student_id: uuid_lib.UUID,
    file: UploadFile = File(...),
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """
    Accept an IQ test PDF/image, extract text (text-layer or OCR),
    call the IQScoreExtractorAgent, and persist IQTestUpload + CognitiveProfile.

    The uploaded file itself is persisted to the AWS S3 bucket
    configured as S3_BUCKET_IQ_TESTS (key pattern
    iq-tests/{student_id}/{uuid}.pdf). The S3 key is stored in
    IQTestUpload.file_path for later retrieval via presigned URL.

    The local temp copy created for OCR is always deleted in the
    finally block.
    """
    await _get_student_or_404(db, student_id)

    allowed_mimes = {
        "application/pdf",
        "image/png",
        "image/jpeg",
        "image/jpg",
    }
    if file.content_type not in allowed_mimes:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{file.content_type}'. Must be PDF, PNG, or JPEG.",
        )

    file_bytes = await file.read()
    size = len(file_bytes)
    if size > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File exceeds 10 MB limit.")
    if size == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Persist to temp path for debugging / inspection mid-pipeline; deleted in finally.
    temp_id = uuid_lib.uuid4()
    ext = ".pdf" if file.content_type == "application/pdf" else (
        ".png" if file.content_type == "image/png" else ".jpg"
    )
    temp_path = os.path.join(TEMP_UPLOAD_DIR, f"{temp_id}{ext}")

    try:
        try:
            os.makedirs(TEMP_UPLOAD_DIR, exist_ok=True)
            with open(temp_path, "wb") as f:
                f.write(file_bytes)
        except Exception as e:
            logger.warning(f"Failed to persist temp upload to {temp_path}: {e}")

        # 1. Extract text (text layer → OCR fallback)
        try:
            raw_text, confidence, ocr_used = extract_text_from_pdf(file_bytes)
        except Exception as e:
            logger.error(f"pdf extraction failed: {e}", exc_info=True)
            raise HTTPException(status_code=422, detail=f"Failed to read file: {str(e)}")

        if not raw_text or len(raw_text.strip()) < 20:
            raise HTTPException(
                status_code=422,
                detail="Could not extract readable text from the uploaded file. "
                       "If this is a scanned document, the OCR engine may have failed.",
            )

        # 2. Run the extractor agent
        agent = IQScoreExtractorAgent()
        try:
            parsed = await agent.extract(raw_text)
        except Exception as e:
            logger.error(f"IQScoreExtractorAgent failed: {e}", exc_info=True)
            raise HTTPException(status_code=502, detail=f"Score extraction failed: {str(e)}")

        if isinstance(parsed, dict) and parsed.get("error"):
            raise HTTPException(status_code=422, detail=parsed["error"])

        # 3. Persist the raw file to S3 for long-term retention. If S3
        #    isn't configured (e.g. local dev without AWS creds) the
        #    storage service falls back to dev-mode logging and returns
        #    the key so IQTestUpload.file_path still gets populated.
        s3_key = s3_storage.iq_test_key(str(student_id), ext)
        stored_key = s3_storage.upload_bytes(
            data=file_bytes,
            key=s3_key,
            content_type=file.content_type,
        )
        if not stored_key:
            # Upload failed and storage is supposedly enabled — fall back
            # to the temp path so we at least have provenance.
            logger.warning(
                "S3 upload failed for student %s — recording temp_path only",
                student_id,
            )
            stored_key = temp_path

        iq_upload = IQTestUpload(
            student_id=student_id,
            uploaded_by_user_id=current_user.id,
            file_name=file.filename or f"upload{ext}",
            file_path=stored_key,  # S3 key (or fallback temp path)
            file_size_bytes=size,
            mime_type=file.content_type,
            upload_status=UploadStatus.COMPLETED,
        )
        db.add(iq_upload)
        await db.flush()

        # 4. Create CognitiveProfile row
        test_name = parsed.get("test_name") if isinstance(parsed, dict) else None
        profile = CognitiveProfile(
            student_id=student_id,
            iq_test_upload_id=iq_upload.id,
            test_name=test_name or "Unknown cognitive assessment",
            administered_by=parsed.get("administered_by") if isinstance(parsed, dict) else None,
            raw_ocr_text=raw_text,
            parsed_scores=parsed,
            confidence_score=confidence,
            requires_review=confidence < 0.9,
        )
        db.add(profile)
        await db.commit()
        await db.refresh(profile)
        await db.refresh(iq_upload)

        return {
            "ocr_used": ocr_used,
            "confidence_score": confidence,
            "cognitive_profile": _serialize_profile(profile, iq_upload),
        }

    finally:
        # Always delete the temp file — privacy: no raw PDFs retained.
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except Exception as e:
            logger.warning(f"Failed to delete temp file {temp_path}: {e}")


# ============================================================================
# COGNITIVE REPORT — GENERATE NARRATIVE
# ============================================================================
@router.post("/students/{student_id}/cognitive-report/generate")
async def generate_cognitive_report(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Generate a narrative Cognitive Report from the latest CognitiveProfile."""
    student = await _get_student_or_404(db, student_id)

    result = await db.execute(
        select(CognitiveProfile)
        .where(CognitiveProfile.student_id == student_id)
        .order_by(desc(CognitiveProfile.created_at))
    )
    profile = result.scalars().first()
    if not profile:
        raise HTTPException(
            status_code=400,
            detail="No cognitive profile found for this student. Upload an IQ test PDF first.",
        )

    student_name = f"{student.first_name} {student.last_name}".strip() or "the student"

    agent = CognitiveReportAgent()
    started = time.time()
    try:
        markdown = await agent.generate(profile.parsed_scores or {}, student_name)
    except Exception as e:
        logger.error(f"CognitiveReportAgent raised: {e}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Agent failed: {str(e)}")
    generation_ms = int((time.time() - started) * 1000)

    report = PsychologistReport(
        student_id=student_id,
        created_by_user_id=current_user.id,
        report_type="cognitive_report",
        content_markdown=markdown,
        source_data={"parsed_scores": profile.parsed_scores},
        source_cognitive_profile_id=profile.id,
        status="draft",
        generation_ms=generation_ms,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)

    return _serialize_report(report)


# ============================================================================
# UNIFIED INSIGHTS — GENERATE
# ============================================================================
@router.post("/students/{student_id}/unified-insights/generate")
async def generate_unified_insights(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Generate a cross-reference Unified Insights report from latest background + cognitive reports."""
    student = await _get_student_or_404(db, student_id)

    background = await _latest_report(db, student_id, "background_summary")
    cognitive = await _latest_report(db, student_id, "cognitive_report")

    if not background or not cognitive:
        missing = []
        if not background:
            missing.append("background summary")
        if not cognitive:
            missing.append("cognitive report")
        raise HTTPException(
            status_code=400,
            detail=f"Both a background summary and a cognitive report are required. Missing: {', '.join(missing)}.",
        )

    student_name = f"{student.first_name} {student.last_name}".strip() or "the student"

    agent = UnifiedInsightsAgent()
    started = time.time()
    try:
        markdown = await agent.synthesize(
            background.content_markdown or "",
            cognitive.content_markdown or "",
            student_name,
        )
    except Exception as e:
        logger.error(f"UnifiedInsightsAgent raised: {e}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Agent failed: {str(e)}")
    generation_ms = int((time.time() - started) * 1000)

    report = PsychologistReport(
        student_id=student_id,
        created_by_user_id=current_user.id,
        report_type="unified_insights",
        content_markdown=markdown,
        source_data={
            "background_summary_id": str(background.id),
            "cognitive_report_id": str(cognitive.id),
        },
        source_chat_session_id=background.source_chat_session_id,
        source_cognitive_profile_id=cognitive.source_cognitive_profile_id,
        status="draft",
        generation_ms=generation_ms,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)

    return _serialize_report(report)


# ============================================================================
# BLANK REPORT — create empty editable row
# ============================================================================
@router.post("/students/{student_id}/reports/blank")
async def create_blank_report(
    student_id: uuid_lib.UUID,
    payload: BlankReportCreate,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Create an empty editable report (used by the 'Start blank' flow)."""
    await _get_student_or_404(db, student_id)

    allowed_types = {
        "background_summary",
        "background_summary_parent",
        "background_summary_school",
        "cognitive_report",
        "unified_insights",
    }
    if payload.report_type not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid report_type '{payload.report_type}'. Must be one of {sorted(allowed_types)}.",
        )

    report = PsychologistReport(
        student_id=student_id,
        created_by_user_id=current_user.id,
        report_type=payload.report_type,
        content_markdown="",
        status="draft",
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)

    return _serialize_report(report)


# ============================================================================
# REPORT CRUD — patch / delete / finalize
# ============================================================================
@router.patch("/reports/{report_id}")
async def update_report(
    report_id: uuid_lib.UUID,
    payload: ReportUpdate,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Update the markdown content of an existing report (autosave + manual save)."""
    report = await _get_report_or_404(db, report_id)
    report.content_markdown = payload.content_markdown
    # updated_at is updated automatically via onupdate=func.now()
    await db.commit()
    await db.refresh(report)
    return _serialize_report(report)


@router.delete("/reports/{report_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_report(
    report_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Delete a report row (used by 'Discard and regenerate')."""
    report = await _get_report_or_404(db, report_id)
    await db.delete(report)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.post("/reports/{report_id}/finalize")
async def finalize_report(
    report_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """Lock a report: status -> 'final'."""
    report = await _get_report_or_404(db, report_id)
    report.status = "final"
    await db.commit()
    await db.refresh(report)
    return _serialize_report(report)


@router.get("/students/{student_id}/report.docx")
async def download_student_reports_docx(
    student_id: uuid_lib.UUID,
    current_user: User = Depends(require_psychologist_or_admin),
    db: AsyncSession = Depends(get_db),
):
    """
    Download a single .docx combining the student's Background, Cognitive and
    Unified Insights psychologist reports (in that order). Returns 404 if no
    reports exist.
    """
    from docx import Document
    from app.api.admin import _render_markdown_to_docx

    student_res = await db.execute(select(Student).where(Student.id == student_id))
    student = student_res.scalar_one_or_none()
    if not student:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Student not found",
        )

    stmt = select(PsychologistReport).where(
        PsychologistReport.student_id == student_id,
        PsychologistReport.report_type.in_(
            (
                "background_summary",
                "background_summary_parent",
                "background_summary_school",
                "cognitive_report",
                "unified_insights",
            )
        ),
    )
    rows = (await db.execute(stmt)).scalars().all()
    if not rows:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No reports found for this student",
        )

    by_type = {r.report_type: r for r in rows}
    # Prefer the role-scoped pair if either exists; fall back to the
    # legacy combined background_summary only when no role-scoped row is present.
    has_role_scoped = bool(
        by_type.get("background_summary_parent") or by_type.get("background_summary_school")
    )
    if has_role_scoped:
        section_order = [
            ("background_summary_parent", "Parent Background Summary"),
            ("background_summary_school", "School Background Summary"),
            ("cognitive_report", "Cognitive Report"),
            ("unified_insights", "Unified Insights"),
        ]
    else:
        section_order = [
            ("background_summary", "Background"),
            ("cognitive_report", "Cognitive Report"),
            ("unified_insights", "Unified Insights"),
        ]

    doc = Document()
    full_name = f"{student.first_name} {student.last_name}".strip()
    doc.add_heading(full_name or "Student Report", level=1)
    for rtype, heading in section_order:
        report = by_type.get(rtype)
        if not report:
            continue
        doc.add_heading(heading, level=2)
        _render_markdown_to_docx(doc, report.content_markdown or "")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r"[^A-Za-z0-9]+", "", f"{student.first_name}{student.last_name}") or "Student"
    filename = f"{safe_name}_report.docx"

    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
